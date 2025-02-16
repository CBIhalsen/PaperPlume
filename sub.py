from docx import Document
from docx.shared import Pt


import yaml
import shutil
from docx.shared import RGBColor

import re
from session_fake import chat
import tiktoken

import random



def load_config():
    with open('./config/prompt.yaml', 'r',encoding='utf-8') as config_file:
        config = yaml.safe_load(config_file)
    return config

# 加载配置
config = load_config()


# 将提纲作为数据返回到sub，sub根据开头空格进行分行提取，通过遍历sub将每一二级标题返回到main

# 创建标题1样式


def create_word(encrypted_name,language):


    source_path ="../documents/word/sources/{}.docx".format(language)
    clone_path ="../documents/word/users/{}{}".format(encrypted_name,".docx")

    # 保存克隆的文档

    shutil.copy2(source_path, clone_path)
    # print("文档克隆完成！")




def nest_outline(outline_list):
    nested_outline = []
    while outline_list:
        level, title = outline_list.pop(0)
        sub_outline = []
        while outline_list and len(outline_list[0][0]) > len(level):
            sub_outline.append(outline_list.pop(0))
        nested_outline.append([title, nest_outline(sub_outline)])
    return nested_outline

# 转换为嵌套列表

class process_outline:
    chat_history = []
    def __init__(self, language=None, style=None, reference=None,title=None, model='gpt-3.5-turbo'):
        self.language = language
        self.model =model
        self.style = style
        self.encoding = tiktoken.encoding_for_model(model)
        self.history= []
        self.title = title
        self.insert_number =18
        self.font_name ='Times New Roman'

    def spilt(self,outline_str,encrypted):
        encrypted_title = encrypted
        # self.chat_history.append(outline_str)
        language = self.language
        if language == 'zh-CN':
            abstract = '摘要'
            self.font_name = 'Songti'
            self.note = config['zh_note']
            self.recommend = config['zh_recommend']

            abstract_input = '请撰写论文摘要，摘要包括关键词,输出语言为中文'
            subheading = "请以小标题 ：{} 。撰写本篇研究论文的一部分。要求保证上下文连贯性和语境一致性。如有需要，可以引入实际案例进行阐述，但需保证案例的完整性和准确性。"
        elif language == 'zh-TW':
            abstract = '摘要'
            self.note = config['zh_tc_note']
            self.recommend = config['zh_tc_recommend']
            self.insert_number = 20
            self.font_name = 'Songti'
            abstract_input = '請撰寫論文摘要，摘要包括關鍵詞，輸出語言爲中文繁体'
            subheading = "請以小標題 ：{} 。撰寫本篇研究論文的一部分。要求保證上下文連貫性和語境一致性。如有需要，可以引入實際案例進行闡述，但需保證案例的完整性和準確性。"
        # elif language == 'en':
        #     abstract = 'Abstract'
        #     abstract_input = 'Please write about {}, you can decide whether to use the actual case. In terms of writing style, you can refer to the style of {} magazine.'
        elif language == 'ja':
            abstract = '抄録'
            self.font_name = 'MS Mincho'
            self.insert_number=12
            self.note = config['ja_note']
            self.recommend = config['ja_recommend']
            abstract_input = 'Please write an abstract of the paper, including keywords. Use Japanese'
            subheading = '小見出し : {} を使用してください。 この研究論文の一部を書いてください。 文脈の一貫性と文脈の整合性が求められます。 必要であれば、実際の事例を紹介して推敲してもよいが、事例の完全性と正確性を確保する必要がある。'
        else:
            abstract = 'Abstract'
            self.insert_number=14
            self.font_name = 'Times New Roman'
            self.note = config['en_note']
            self.recommend = config['en_recommend']
            abstract_input = 'Please write an abstract of the paper, including keywords'
            subheading = 'Please use the subheading : {} . Write a portion of this research paper. Contextual coherence and contextual consistency are required. If necessary, actual cases can be introduced for elaboration, but the completeness and accuracy of the cases need to be ensured.'




        result = replace_spaces_with_hashes(outline_str)
        # print("result",result)
        lines = result.split('\n')
        lines = [line for line in lines if line.strip() != '']

        depths_and_titles = [get_depth_and_title(line) for line in lines]

        chapters = []
        current_chapter = []

        for depth, title in depths_and_titles:
            if depth == 2:
                if current_chapter:
                    chapters.append(current_chapter)
                current_chapter = [(depth, title)]
            else:
                current_chapter.append((depth, title))

        chapters.append(current_chapter)
        a = 0
        outline_str =  self.note + self.title + "\n" + outline_str
        messages = [{"role": "system", "content": self.recommend}]  # 添加第一次的提示词到历史记录
        messages.append({"role": "user", "content": outline_str})
        total_tokens = 0
        for chapter in chapters:
            deepest_depth = max(section[0] for section in chapter)
            for depth, title in chapter:
                # print(title)  # 打印出该章节所有层的标题
                max_history = 9
                min_history = 7
                # print(len(history))




                if len(messages) > max_history:  # 历史长度最大为9
                    messages = [messages[0],messages[1]] + messages[-min_history:]




                # for message in self.chat_history:
                #     messages.append({"role": "assistant", "content": message})



                # 检查总 token 数是否超出限制
                request_tokens = sum(len(self.encoding.encode(message['content'])) for message in messages)
                # print(total_tokens)

                # print("history:",history)
                max_tokens = 2650

                if request_tokens > max_tokens:
                    print('当前tokens:', request_tokens)
                    print('当前长度', len(messages))
                    # print("当前token数量:", request_tokens, "开始截断")
                    i = 0
                    while request_tokens > max_tokens:
                        i = i + 1
                        if len(messages) > 5:
                            del messages[-5]
                        elif len(messages) == 5:
                            del messages[-4]
                        else:
                            del messages[-2]
                        request_tokens = sum(len(self.encoding.encode(message['content'])) for message in messages)

                if depth != deepest_depth:  # 如果当前标题不是最深层的标题
                    if depth !=2:
                        if depth-2 == deepest_depth-depth:
                            if depth >2:
                                depth=3
                        elif depth-2 < deepest_depth-depth:
                            depth=3
                        else:
                            depth=4
                    a += 1
                    if a == 5:
                        write_abstract(abstract, self.insert_number, encrypted_title, option2="head1")
                        model_reply, tokens = chat(self.language, self.style, model=self.model).chat_with_model(
                            prompt=abstract_input, title=self.title, messages=messages)
                        total_tokens = total_tokens + tokens + request_tokens

                        #     break  # 如果没有生错误，则跳出循环
                        # except Exception as e:
                        #     print(f"Error occurred: {e}")
                        #     print("Retrying...")

                        # self.chat_history.append(model_reply)
                        messages.append({"role": "user", "content": model_reply})
                        messages[2], messages[-1] = messages[-1], messages[2]
                        # model_reply = indent_string(model_reply)
                        # model_reply = randomize_spaces(model_reply)
                        model_reply = correct_grammarly(language=self.language).correct(model_reply)



                        input_abstract = "{}{}{}".format(abstract, "\n", model_reply)
                        write_abstract(input_abstract, self.insert_number+1, encrypted_title, option3="body")



                    # print("Title",title)
                    generate_word(data=title, encrypted_title=encrypted_title, style=depth,font_name=self.font_name)
                    # print("非深层标题等级", depth)
                else:
                    if depth >4:
                        depth=4

                    a += 1
                    if a == 1:
                        continue
                    if a == 5:
                        write_abstract(abstract, self.insert_number, encrypted_title, option2="head1")
                        model_reply, tokens = chat(self.language, self.style, model=self.model).chat_with_model(
                            prompt=abstract_input, title=self.title, messages=messages)
                        total_tokens = total_tokens + tokens +request_tokens

                        #     break  # 如果没有生错误，则跳出循环
                        # except Exception as e:
                        #     print(f"Error occurred: {e}")
                        #     print("Retrying...")

                        # self.chat_history.append(model_reply)
                        messages.append({"role": "user", "content": model_reply})
                        messages[2], messages[-1] = messages[-1], messages[2]
                        # model_reply = indent_string(model_reply)
                        # model_reply = randomize_spaces(model_reply)
                        model_reply = correct_grammarly(language=self.language).correct(model_reply)
                        input_abstract = "{}{}{}".format(abstract, "\n", model_reply)
                        write_abstract(input_abstract, self.insert_number, encrypted_title, option3="body")



                    # print(title, "sssss")
                    user_input = title
                    print("写作小标题", user_input,"添加标题等级",depth)
                    generate_word(data=title, encrypted_title=encrypted_title, style=depth,font_name=self.font_name)

                    upload_input = subheading.format(user_input, self.style)
                    model_reply, tokens = chat(self.language, self.style, model=self.model).chat_with_model(prompt=upload_input, title=self.title, messages=messages)
                    total_tokens = total_tokens + tokens + request_tokens

                    # self.chat_history.append(model_reply)
                    messages.append({"role": "user", "content": model_reply})
                    print(model_reply)
                    model_reply = correct_grammarly(language=self.language).correct(model_reply)
                    # model_reply = indent_string(model_reply)2

                    generate_word(data=model_reply, encrypted_title=encrypted_title, style=70,font_name=self.font_name)
        return total_tokens

class correct_grammarly():

    def __init__(self, language=None):
        self.language = language
    def correct(self,text):
        if self.language == 'en':
            return self.normalize_text_en(text)
        else:
            return self.normalize_text(text)

    def normalize_text_en(self,text):
        # 在句子分隔符和后面的单词之间添加一个空格
        text = re.sub(r'([.,;:!?"])(\w)', r'\1 \2', text)

        # 在句号前不添加空格
        text = re.sub(r'(\w) \.', r'\1.', text)

        # 移除多余的空格（保留换行符）
        text = re.sub(r'(?<!\n) +', ' ', text)
        text = re.sub(r'\n\s*\n', '\n', text)

        # 在每个段落前添加两个空格
        text = re.sub(r'\n', '\n  ', text)
        text = re.sub(r'\s*\d+\s*$', '', text)

        return text

    def normalize_text(self,text):
        # 在句子分隔符后面的单词之间移除空格
        text = re.sub(r'([。,，;；:：!?！？"“”])(\w)', r'\1\2', text)

        # 移除多余的空格（保留换行符）
        text = re.sub(r'(?<!\n) +', ' ', text)

        # 删除空行
        text = re.sub(r'\n\s*\n', '\n', text)


        # 在每个段落前添加两个空格
        text = re.sub(r'\n', '\n  ', text)
        text = re.sub(r'\s*\d+\s*$', '', text)

        return text



def randomize_spaces(text):
    result = ''
    for char in text:
        if char == ' ':
            if random.random() < 0.7:  # 70%的概率变成半角空格
                result += ' '
            else:  # 30%的概率变成全角空格
                result += '　'
        else:
            result += char
    return result

# 示例输入
# input_text = acc

# 随机变换空格
# randomized_text = randomize_spaces(input_text)
#
# # 输出结果
# print(randomized_text)


def get_depth_and_title(line):
    match = re.match(r'#+', line)
    if match is None:
        depth = 0
    else:
        depth = len(match.group())
    title = line.lstrip('# ').rstrip()
    return depth, title



def replace_spaces_with_hashes(string):
    lines = string.split('\n')
    result = []
    for line in lines:

        num_spaces = len(line) - len(line.lstrip())
        if line.strip() != ''  and num_spaces ==0:
            result.append('#' * 2 + line.lstrip())
        if line.strip() != '' and num_spaces !=0:
            result.append('#' * (num_spaces+1) + line.lstrip())
    return '\n'.join(result)




def print_outline(nested_outline, level=0, count=[0]):
    for title, sub_outline in nested_outline:
        # print('  ' * level + title)
        count[0] += 1
        print_outline(sub_outline, level + 1, count)
    return count[0]

# print_count = print_outline(nested_outline)
# print("Total print count: ", print_count)
def count_indent_width(s):
    width = 0
    for c in s:
        if c == ' ':
            width += 1
        elif c == '　':  # 这是一个全角空格
            width += 2
        else:
            break
    return width



# def pprocess_outline(outline_str, encryption_str):
#     outline_list = outline_str.split('\n')
#     outline_list = [line.strip() for line in outline_list if line.strip() != '']
#
#     outline_hierarchy = []
#
#     for line in outline_list:
#         match = re.match(r'^(\s*)(\d+\.\d+)*(.*)$', line)
#         indent, numbering, title = match.groups()
#
#         level = len(indent) // 3 + 1
#
#         outline_hierarchy.append((title, level))
#
#     max_level = max(outline_hierarchy, key=lambda x: x[1])[1]
#
#     for title, level in outline_hierarchy:
#         print('  ' * (level - 1) + title)
#     print(max_level,'shens')
#     return max_level




def generate_word(data,encrypted_title,style,font_name):
    # 处理可选关键字参数
    if style == 2:
        doc = Document("../documents/word/users/"+encrypted_title +".docx")

        doc.styles['Heading 1'].font.name = font_name
        #创建标题1样式
        title1_style = doc.styles['Heading 1']

        title1_style.paragraph_format.space_before = Pt(12)
        title1_style.paragraph_format.space_after = Pt(6)
        title1_style.paragraph_format.keep_with_next = True
        title1_style.font.size = Pt(14)
        title1_style.font.bold = True
        title1_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 设置为黑色
        paragraph = doc.add_paragraph()

        # 设置段落样式为标题1
        paragraph.style = title1_style
        # print(title1_style)
        # print(paragraph)
        # 添加文本为标题1
        paragraph.add_run(data)
        # print("选项1:", option1)
        doc.save( "../documents/word/users/"+ encrypted_title + '.docx')
    elif style ==3:
        doc = Document("../documents/word/users/"+ encrypted_title + ".docx")

        # 创建标题2样式
        doc.styles['Heading 2'].font.name = font_name
        title2_style = doc.styles['Heading 2']

        title2_style.paragraph_format.space_before = Pt(12)
        title2_style.paragraph_format.space_after = Pt(6)
        title2_style.paragraph_format.keep_with_next = True
        title2_style.font.size = Pt(14)
        title2_style.font.bold = True
        title2_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 设置为黑色

        paragraph = doc.add_paragraph()

        # print(title2_style)


        paragraph.style = title2_style

        paragraph.add_run(data)
        doc.save("../documents/word/users/"+encrypted_title + '.docx')

    elif style ==4:
        doc = Document("../documents/word/users/"+ encrypted_title + ".docx")
        doc.styles['Heading 3'].font.name = font_name
        # 创建标题2样式
        title2_style = doc.styles['Heading 3']

        title2_style.paragraph_format.space_before = Pt(12)
        title2_style.paragraph_format.space_after = Pt(6)
        title2_style.paragraph_format.keep_with_next = True
        title2_style.font.size = Pt(14)
        title2_style.font.bold = True
        title2_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 设置为黑色

        paragraph = doc.add_paragraph()

        # print(title2_style)


        paragraph.style = title2_style

        paragraph.add_run(data)
        doc.save("../documents/word/users/"+encrypted_title + '.docx')

    elif style ==5:
        doc = Document("../documents/word/users/"+ encrypted_title + ".docx")
        doc.styles['Heading 4'].font.name = font_name
        # 创建标题2样式
        title2_style = doc.styles['Heading 4']

        title2_style.paragraph_format.space_before = Pt(12)
        title2_style.paragraph_format.space_after = Pt(6)
        title2_style.paragraph_format.keep_with_next = True
        title2_style.font.size = Pt(14)
        title2_style.font.bold = True
        title2_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 设置为黑色

        paragraph = doc.add_paragraph()

        # print(title2_style)


        paragraph.style = title2_style

        paragraph.add_run(data)
        doc.save("../documents/word/users/"+encrypted_title + '.docx')
    else:
        doc = Document("../documents/word/users/"+encrypted_title + ".docx")
        doc.styles['Normal'].font.size = Pt(12)
        doc.styles['Normal'].font.name = font_name
        # 创建正文样式
        body_style = doc.styles['Normal']
        body_style.paragraph_format.space_before = Pt(0)
        body_style.paragraph_format.space_after = Pt(0)
        body_style.paragraph_format.line_spacing = 1.2


        paragraph = doc.add_paragraph()
        # option3 = kwargs['option3']
        # print(body_style)

        # 设置段落样式为标题1
        # paragraph.style = body_style


        # 为每一行创建一个新的段落

        for line in data.split("\n"):
            if line.strip():  # 这将过滤掉空字符串
                paragraph = doc.add_paragraph()
                paragraph.style = body_style
                line = indent_string(line)
                paragraph.add_run(line)
        # lines = data.split("\n")
        # for line in lines:
        #     paragraph = doc.add_paragraph()
        #     # 设置段落样式为标题1
        #     paragraph.style = body_style
        #     line = indent_string(line)
        #
        #     # 添加文本为标题1
        #     paragraph.add_run(line)



        # paragraph.add_run(data)
        # for style in doc.styles:
        #     print(style.name)
        doc.save("../documents/word/users/"+encrypted_title + '.docx')










def write_abstract(data,number,encrypted_title,**kwargs):
    if 'option2' in kwargs:
        doc = Document("../documents/word/users/"+ encrypted_title + ".docx")
        doc.styles['Normal'].font.name = '宋体'
        # 创建标题2样式
        title2_style = doc.styles['Heading 2']

        title2_style.paragraph_format.space_before = Pt(12)
        title2_style.paragraph_format.space_after = Pt(6)
        title2_style.paragraph_format.keep_with_next = True
        title2_style.font.size = Pt(14)
        title2_style.font.bold = True
        title2_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 设置为黑色
        title2_style.paragraph_format.alignment = 0

        paragraph = doc.add_paragraph()
        option2 = kwargs['option2']
        # print(title2_style)

        paragraph.style = title2_style
        paragraph = doc.paragraphs[number]
        paragraph.add_run(data)

        doc.save("../documents/word/users/"+encrypted_title + '.docx')

    if 'option3' in kwargs:



        doc = Document("../documents/word/users/" + encrypted_title + ".docx")
        doc.styles['Normal'].font.name = '宋体'
            # 创建正文样式
        body_style = doc.styles['Normal']
        body_style.paragraph_format.space_before = Pt(6)
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.line_spacing = 1.5
        body_style.paragraph_format.alignment = 0

        paragraph = doc.add_paragraph()

        # print(body_style)

        # 设置段落样式为标题1
        paragraph.style = body_style


        # 添加文本为标题1
            # print(paragraph.style)
        paragraph = doc.paragraphs[number]
        paragraph.add_run(data)
            # for style in doc.styles:
            #     print(style.name)
        doc.save("../documents/word/users/" + encrypted_title + '.docx')


def write_content(data,number,encrypted_title,heading):
    doc = Document("../documents/word/users/" + encrypted_title + ".docx")
    doc.styles['Normal'].font.name = '宋体'

    # 创建标题1样式
    title1_style = doc.styles[heading]

    title1_style.paragraph_format.space_before = Pt(12)
    title1_style.paragraph_format.space_after = Pt(6)
    title1_style.paragraph_format.keep_with_next = True
    title1_style.font.size = Pt(14)
    title1_style.font.bold = True
    title1_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 设置为黑色
    paragraph = doc.add_paragraph()

    # 设置段落样式为标题1
    paragraph.style = title1_style
    # print(title1_style)
    # print(paragraph)
    # 添加文本为标题1
    paragraph = doc.paragraphs[number]
    paragraph.add_run(data)
    # print("选项1:", option1)
    doc.save("../documents/word/users/" + encrypted_title + '.docx')

def indent_string(input_string):
    lines = input_string.split('\n')
    indented_lines = ['  ' + line for line in lines]
    return '\n'.join(indented_lines)





def randomize_spaces(text):
    result = ''
    for char in text:
        if char == ' ':
            if random.random() < 0.5:  # 70%的概率变成半角空格
                result += ' '
            else:  # 30%的概率变成全角空格
                result += '　'
        else:
            result += char
    return result

def process_text(text):
    sentences = re.split(r'(?<=[.!?]) +', text)
    for i in range(len(sentences)):
        words = sentences[i].split()
        if len(words) > 30:
            sentences[i] = randomize_spaces(sentences[i])
    return ' '.join(sentences)

# 测试代码


def normalize_text_en(text):
    # 在句子分隔符和后面的单词之间添加一个空格
    text = re.sub(r'([.,;:!?"])(\w)', r'\1 \2', text)

    # 在句号前不添加空格
    text = re.sub(r'(\w) \.', r'\1.', text)

    # 移除多余的空格（保留换行符）
    text = re.sub(r'(?<!\n) +', ' ', text)

    # 在每个段落前添加两个空格
    text = re.sub(r'\n', '\n  ', text)

    return text


def normalize_text(text):
    # 在句子分隔符后面的单词之间移除空格
    text = re.sub(r'([。,，;；:：!?！？"“”])(\w)', r'\1\2', text)

    # 移除多余的空格（保留换行符）
    text = re.sub(r'(?<!\n) +', ' ', text)

    # 在每个段落前添加两个空格
    text = re.sub(r'\n', '\n  ', text)

    return text